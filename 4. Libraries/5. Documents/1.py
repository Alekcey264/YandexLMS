from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT
from sys import stdin
from docx.shared import Pt, Inches


#text = list(map(str.strip, stdin))
text = "Dear Kate and Nick,\nWe are looking forward very much to your visit to our country this summer. We are expecting you at the beginning of July and are hoping that you may stay until the end of the month or longer.\nWe consider it a privilege for us to receive you as guests in our house. We are very grateful indeed to you for consenting to come and stay with us. We are looking forward to oﬀering you hospitality in return for the hospitality you have kindly given us on many occasions.\nWe want you to understand that we will see to all your needs while you are with us and to any expenses that may arise.\nYours sincerely,\nJohn and Mary".split('\n')

greenting = text[0]
closing = text[-2:]
body = text[1:-2]

doc = Document()
p = doc.add_paragraph(greenting)
p.alignment = WD_PARAGRAPH_ALIGNMENT.RIGHT
run = p.runs[0]
run.font.name = "Arial"
run.italic = True
run.font.size = Pt(11)

for line in body:
    p = doc.add_paragraph(line)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.JUSTIFY
    p.paragraph_format.first_line_indent = Inches(0.5)
    run = p.runs[0]
    run.font.name = "Times New Roman"
    run.size = Pt(12)

for line in closing[:-1]:
    p = doc.add_paragraph(line)
    p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

p = doc.add_paragraph(closing[-1])
p.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
run = p.runs[0]
run.bold = True

doc.save('D:\\Yandex LMS\\4. Libraries\\5. Documents\\letter.docx')