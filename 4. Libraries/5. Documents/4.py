from docxtpl import DocxTemplate
from docx import Document
from docx.shared import Inches
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT


def template(f_name):
    doc = Document()
    p = doc.add_heading("Самостоятельная работа № {{ work }}", level=1)
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
    p = doc.add_heading("Вариант № {{ variant }}", level=2)
    p.paragraph_format.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    doc.add_paragraph("Задание 1\n{{ task1 }}", style="List Number")
    doc.add_paragraph("Задание 2\n{{ task2 }}", style="List Number")
    doc.add_paragraph("Задание 3\n{{ task3 }}", style="List Number")
    p = doc.add_paragraph("Выполнил: {{ name }}")
    p.paragraph_format.first_line_indent = Inches(0.5)
    run = p.runs[0]
    run.italic = True
    doc.save(f_name)


def prepare_doc(f_in_name, **dic):
    doc = DocxTemplate(f_in_name)
    doc.render(dic)
    doc.save("D:\\Yandex LMS\\4. Libraries\\5. Documents\\res.docx")


template("D:\\Yandex LMS\\4. Libraries\\5. Documents\\tpl.docx")
dic = {"work": 1,
       "variant": 3,
       "task1": "2 + 3 =",
       "task2": "12 * 4 =",
       "task3": "How many zeros in binary digit of 8?",
       "name": "John Hopkins"}
prepare_doc("D:\\Yandex LMS\\4. Libraries\\5. Documents\\tpl.docx", **dic)