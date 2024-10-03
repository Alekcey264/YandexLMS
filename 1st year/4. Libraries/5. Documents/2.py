from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT

#text = list(map(str.strip, stdin))
text = "WBC\nRBC\nHGB\nHTC\nMCV\nMCH\nMCHC\nRDW\nPLT\nMPV\nPTC\nNEU\nLYMP\nMONO\nEO\nBaso".split("\n")

doc = Document()

title = doc.add_heading(level=0)
title_run = title.add_run("Blood test")
title.alignment = WD_PARAGRAPH_ALIGNMENT.CENTER
title.runs[0].bold = True

table = doc.add_table(rows=len(text) + 1, cols=3, style="Table Grid")
cells = table.rows[0].cells
cells[0].text = "indicator"
cells[1].text = "norm"
cells[2].text = "value"


for cell in cells:
    cell.paragraphs[0].runs[0].bold = True
    cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.CENTER


for i, indicator in enumerate(text, start=1):
    cu_cell = table.cell(i, 0)
    cu_cell.text = indicator
    cu_cell.paragraphs[0].alignment = WD_PARAGRAPH_ALIGNMENT.LEFT

doc.save('D:\\Yandex LMS\\4. Libraries\\5. Documents\\analysis.docx')