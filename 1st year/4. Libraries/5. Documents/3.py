import sys
from docx import Document
from docx.shared import Inches

text = list(map(str.rstrip, sys.stdin))
doc = Document()
doc.add_heading(text[0], 0)
for item in text:
    index_of = item.count(" ") - len(item.split()) + 1
    if (index_of + 2) // 2 < 2:
        set_style = 'List Bullet'
    else:
        set_style = 'List Bullet ' + str((index_of + 2) // 2)
    p = doc.add_paragraph(item.strip(), style=set_style)
    p.paragraph_format.left_indent = Inches((index_of // 2) * 0.25)

doc.save('D:\\Yandex LMS\\4. Libraries\\5. Documents\\systematization.docx')